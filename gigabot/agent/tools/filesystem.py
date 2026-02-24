"""File-system tools: file operations and project management."""

import difflib
import mimetypes
import shutil
from pathlib import Path
from typing import Any

from gigabot.agent.tools.base import Tool


# ---------------------------------------------------------------------------
# Helper readers / writers
# ---------------------------------------------------------------------------

def _read_pdf(file_path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        import pdfplumber
    except ImportError:
        return "Error: Reading PDF requires the pdfplumber package. Install with: pip install pdfplumber"
    try:
        with pdfplumber.open(file_path) as pdf:
            parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return "\n\n".join(parts) if parts else "The PDF file contains no extractable text (may be scanned/image-only)."
    except Exception as e:
        return f"Error reading PDF: {e}"


def _read_docx(file_path: Path) -> str:
    """Extract text and table content from a DOCX file."""
    try:
        import docx
    except ImportError:
        return "Error: python-docx not installed"
    try:
        doc = docx.Document(str(file_path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        if not parts:
            return "The DOCX file contains no text."
        return "\n".join(parts)
    except Exception as e:
        return f"Error reading DOCX: {e}"


def _read_excel(file_path: Path) -> str:
    """Extract data from an Excel file (.xlsx / .xls)."""
    try:
        import openpyxl
    except ImportError:
        return "Error: Reading Excel requires the openpyxl package. Install with: pip install openpyxl"
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if rows:
                parts.append(f"Sheet: {sheet.title}")
                for row in rows:
                    parts.append("\t".join(str(c) if c is not None else "" for c in row))
        wb.close()
        return "\n".join(parts) if parts else "The Excel file contains no data."
    except Exception as e:
        return f"Error reading Excel: {e}"


def _read_text_with_encoding(file_path: Path) -> str:
    """Read a text file, trying UTF-8 then common fallbacks."""
    raw = file_path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _smart_read(file_path: Path) -> str:
    """Dispatch to the correct reader based on file extension."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(file_path)
    if ext in (".docx", ".doc"):
        return _read_docx(file_path)
    if ext in (".xlsx", ".xls"):
        return _read_excel(file_path)
    mime = mimetypes.guess_type(str(file_path))[0] or ""
    if mime.startswith("image/"):
        return f"This is an image file ({file_path.suffix}). Use other tools to process images."
    return _read_text_with_encoding(file_path)


def _write_docx(file_path: Path, content: str) -> str:
    """Create a DOCX file with the given text content."""
    try:
        import docx
    except ImportError:
        return "Error: Creating DOCX requires python-docx. Install with: pip install python-docx"
    doc = docx.Document()
    for block in content.strip().split("\n\n"):
        doc.add_paragraph(block.replace("\n", " "))
    doc.save(str(file_path))
    return f"Successfully wrote DOCX to {file_path}"


def _write_xlsx(file_path: Path, content: str) -> str:
    """Create an XLSX file.  Rows split by newline, columns by tab."""
    try:
        import openpyxl
    except ImportError:
        return "Error: Creating Excel requires openpyxl. Install with: pip install openpyxl"
    wb = openpyxl.Workbook()
    ws = wb.active
    for row_idx, line in enumerate(content.strip().splitlines(), start=1):
        cells = line.split("\t")
        for col_idx, val in enumerate(cells, start=1):
            ws.cell(row=row_idx, column=col_idx, value=val.strip() if isinstance(val, str) else val)
    wb.save(str(file_path))
    return f"Successfully wrote Excel to {file_path}"


# ---------------------------------------------------------------------------
# Path resolution (importable by other modules, e.g. message.py)
# ---------------------------------------------------------------------------

def _resolve_path(
    path_str: str,
    workspace: Path | None = None,
    allowed_dir: Path | None = None,
) -> Path:
    """Resolve *path_str* against *workspace* (when relative) and enforce *allowed_dir*."""
    p = Path(path_str).expanduser()
    if not p.is_absolute() and workspace:
        p = workspace / p
    resolved = p.resolve()
    if allowed_dir and not str(resolved).startswith(str(allowed_dir.resolve())):
        raise PermissionError(f"Path {path_str} is outside allowed directory {allowed_dir}")
    return resolved


# ---------------------------------------------------------------------------
# Tool classes
# ---------------------------------------------------------------------------

class FileTool(Tool):
    """Unified file operations: read, write, edit, list, move."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None) -> None:
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "file"

    @property
    def description(self) -> str:
        return (
            "Работа с файлами: чтение (txt/pdf/docx/xlsx), создание, "
            "редактирование, просмотр каталогов, перемещение"
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["read", "write", "edit", "list", "move"],
                    "description": "Action to perform",
                },
                "path": {
                    "type": "string",
                    "description": "File or directory path",
                },
                "content": {
                    "type": "string",
                    "description": "Content for write action",
                },
                "old_text": {
                    "type": "string",
                    "description": "Text to find (edit action)",
                },
                "new_text": {
                    "type": "string",
                    "description": "Replacement text (edit action)",
                },
                "destination": {
                    "type": "string",
                    "description": "Target path for move action",
                },
            },
            "required": ["action"],
        }

    async def execute(self, action: str, **kwargs: Any) -> str:
        dispatch = {
            "read": self._read,
            "write": self._write,
            "edit": self._edit,
            "list": self._list,
            "move": self._move,
        }
        handler = dispatch.get(action)
        if not handler:
            return f"Error: unknown action '{action}'. Use: read, write, edit, list, move"
        try:
            return await handler(**kwargs)
        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error ({action}): {e}"

    async def _read(self, path: str = "", **_: Any) -> str:
        if not path:
            return "Error: 'path' is required for read. Пример: file(action='read', path='documents/file.txt')"
        file_path = _resolve_path(path, self._workspace, self._allowed_dir)
        if not file_path.exists():
            return f"Error: File not found: {path}"
        if not file_path.is_file():
            return f"Error: Not a file: {path}"
        return _smart_read(file_path)

    async def _write(self, path: str = "", content: str = "", **_: Any) -> str:
        if not path:
            return "Error: 'path' is required for write. Пример: file(action='write', path='output.txt', content='...')"
        if not content:
            return "Error: 'content' is required for write. Пример: file(action='write', path='output.txt', content='текст файла')"
        file_path = _resolve_path(path, self._workspace, self._allowed_dir)
        file_path.parent.mkdir(parents=True, exist_ok=True)
        ext = file_path.suffix.lower()
        if ext == ".docx":
            return _write_docx(file_path, content)
        if ext == ".xlsx":
            return _write_xlsx(file_path, content)
        file_path.write_text(content, encoding="utf-8")
        return f"Successfully wrote {len(content)} bytes to {file_path}"

    async def _edit(self, path: str = "", old_text: str = "", new_text: str = "", **_: Any) -> str:
        if not path:
            return "Error: 'path' is required for edit"
        if not old_text:
            return "Error: 'old_text' is required for edit"
        file_path = _resolve_path(path, self._workspace, self._allowed_dir)
        if not file_path.exists():
            return f"Error: File not found: {path}"

        content = file_path.read_text(encoding="utf-8")
        if old_text not in content:
            return self._not_found_message(old_text, content, path)

        count = content.count(old_text)
        if count > 1:
            return (
                f"Warning: old_text appears {count} times. "
                "Please provide more context to make it unique."
            )

        new_content = content.replace(old_text, new_text, 1)
        file_path.write_text(new_content, encoding="utf-8")
        return f"Successfully edited {file_path}"

    async def _list(self, path: str = "", **_: Any) -> str:
        if not path:
            return "Error: 'path' is required for list"
        dir_path = _resolve_path(path, self._workspace, self._allowed_dir)
        if not dir_path.exists():
            return f"Error: Directory not found: {path}"
        if not dir_path.is_dir():
            return f"Error: Not a directory: {path}"

        items: list[str] = []
        for item in sorted(dir_path.iterdir()):
            prefix = "\U0001f4c1 " if item.is_dir() else "\U0001f4c4 "
            items.append(f"{prefix}{item.name}")
        return "\n".join(items) if items else f"Directory {path} is empty"

    async def _move(self, path: str = "", destination: str = "", **_: Any) -> str:
        if not path:
            return "Error: 'path' is required for move. Пример: file(action='move', path='file.txt', destination='folder/file.txt')"
        if not destination:
            return "Error: 'destination' is required for move. Пример: file(action='move', path='file.txt', destination='folder/file.txt')"
        src = _resolve_path(path, self._workspace, self._allowed_dir)
        dst = _resolve_path(destination, self._workspace, self._allowed_dir)
        if not src.exists():
            return f"Error: Source not found: {path}"
        if not src.is_file():
            return f"Error: Source is not a file: {path}"
        if dst.is_dir():
            dst = dst / src.name
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dst))
        return f"Moved {src} \u2192 {dst}"

    @staticmethod
    def _not_found_message(old_text: str, content: str, path: str) -> str:
        lines = content.splitlines(keepends=True)
        old_lines = old_text.splitlines(keepends=True)
        window = len(old_lines)

        best_ratio, best_start = 0.0, 0
        for i in range(max(1, len(lines) - window + 1)):
            ratio = difflib.SequenceMatcher(
                None, old_lines, lines[i : i + window],
            ).ratio()
            if ratio > best_ratio:
                best_ratio, best_start = ratio, i

        if best_ratio > 0.5:
            diff = "\n".join(difflib.unified_diff(
                old_lines,
                lines[best_start : best_start + window],
                fromfile="old_text (provided)",
                tofile=f"{path} (actual, line {best_start + 1})",
                lineterm="",
            ))
            return (
                f"Error: old_text not found in {path}.\n"
                f"Best match ({best_ratio:.0%} similar) at line {best_start + 1}:\n{diff}"
            )
        return f"Error: old_text not found in {path}. No similar text found. Verify the file content."


class ProjectTool(Tool):
    """Project folder management: create, list, add/delete subfolders."""

    _SUBFOLDERS = (
        "\u0414\u043e\u0433\u043e\u0432\u043e\u0440\u044b",
        "\u0414\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430\u0446\u0438\u044f",
        "\u0421\u043c\u0435\u0442\u044b",
        "\u0424\u043e\u0442\u043e",
        "\u041f\u0435\u0440\u0435\u043f\u0438\u0441\u043a\u0430",
    )

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None) -> None:
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "project"

    @property
    def description(self) -> str:
        return (
            "Управление папками проектов: создание, просмотр, подпапки, "
            "перемещение файлов в проект"
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["create", "list", "add_folder", "delete_folder", "move_file"],
                    "description": "Action to perform",
                },
                "name": {
                    "type": "string",
                    "description": "Project name",
                },
                "folder_name": {
                    "type": "string",
                    "description": "Subfolder name (for add_folder / delete_folder / move_file)",
                },
                "file_path": {
                    "type": "string",
                    "description": "Path to file to move into project (for move_file)",
                },
            },
            "required": ["action"],
        }

    async def execute(self, action: str, **kwargs: Any) -> str:
        dispatch = {
            "create": self._create,
            "list": self._list,
            "add_folder": self._add_folder,
            "delete_folder": self._delete_folder,
            "move_file": self._move_file,
        }
        handler = dispatch.get(action)
        if not handler:
            return f"Error: unknown action '{action}'. Use: create, list, add_folder, delete_folder, move_file"
        try:
            return await handler(**kwargs)
        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error ({action}): {e}"

    @property
    def _projects_dir(self) -> Path:
        if not self._workspace:
            raise RuntimeError("workspace is not configured")
        return self._workspace / "projects"

    async def _create(self, name: str = "", **_: Any) -> str:
        if not name:
            return "Error: 'name' is required for create"
        project_dir = self._projects_dir / name
        if project_dir.exists():
            return f"Error: Project '{name}' already exists at {project_dir}"

        project_dir.mkdir(parents=True, exist_ok=True)
        created: list[str] = []
        for sub in self._SUBFOLDERS:
            (project_dir / sub).mkdir(exist_ok=True)
            created.append(sub)

        structure = "\n".join(f"  \U0001f4c1 {s}" for s in created)
        return f"Project '{name}' created at {project_dir}\nStructure:\n{structure}"

    async def _list(self, **_: Any) -> str:
        projects_dir = self._projects_dir
        if not projects_dir.exists():
            return "No projects directory found."
        dirs = sorted(p.name for p in projects_dir.iterdir() if p.is_dir())
        if not dirs:
            return "No projects found."
        lines = [f"\U0001f4c1 {d}" for d in dirs]
        return f"Projects ({len(dirs)}):\n" + "\n".join(lines)

    async def _add_folder(self, name: str = "", folder_name: str = "", **_: Any) -> str:
        if not name:
            return "Error: 'name' (project name) is required"
        if not folder_name:
            return "Error: 'folder_name' is required"
        project_dir = self._projects_dir / name
        if not project_dir.exists():
            return f"Error: Project '{name}' not found"
        new_folder = project_dir / folder_name
        if new_folder.exists():
            return f"Folder '{folder_name}' already exists in project '{name}'"
        new_folder.mkdir(parents=True, exist_ok=True)
        return f"Added folder '{folder_name}' to project '{name}'"

    async def _delete_folder(self, name: str = "", folder_name: str = "", **_: Any) -> str:
        if not name:
            return "Error: 'name' (project name) is required"
        if not folder_name:
            return "Error: 'folder_name' is required"
        project_dir = self._projects_dir / name
        if not project_dir.exists():
            return f"Error: Project '{name}' not found"
        target = project_dir / folder_name
        if not target.exists():
            return f"Error: Folder '{folder_name}' not found in project '{name}'"
        shutil.rmtree(target)
        return f"Deleted folder '{folder_name}' from project '{name}'"

    @staticmethod
    def _clean_telegram_filename(name: str) -> str:
        """Strip Telegram file_id prefix like 'BQACAgIAAxkBAAIC_05_' from filename."""
        import re
        cleaned = re.sub(r'^[A-Za-z]{2}[A-Za-z0-9]{10,}_\d+_', '', name)
        return cleaned if cleaned and cleaned != name else name

    async def _move_file(self, name: str = "", folder_name: str = "", file_path: str = "", **_: Any) -> str:
        if not name:
            return (
                "Error: 'name' (project name) is required. "
                "Пример: project(action='move_file', name='Коттедж', folder_name='Договора', file_path='/path/to/file.pdf')"
            )
        if not file_path:
            return (
                "Error: 'file_path' is required. "
                "Пример: project(action='move_file', name='Коттедж', folder_name='Договора', file_path='/path/to/file.pdf')"
            )
        project_dir = self._projects_dir / name
        if not project_dir.exists():
            available = sorted(p.name for p in self._projects_dir.iterdir() if p.is_dir()) if self._projects_dir.exists() else []
            hint = f" Доступные проекты: {', '.join(available)}" if available else ""
            return f"Error: Project '{name}' not found.{hint}"

        src = _resolve_path(file_path, self._workspace, self._allowed_dir)
        if not src.exists():
            return f"Error: File not found: {file_path}"
        if not src.is_file():
            return f"Error: Not a file: {file_path}"

        subfolders = sorted(d.name for d in project_dir.iterdir() if d.is_dir())

        if folder_name:
            target_dir = project_dir / folder_name
            if not target_dir.exists():
                hint = f" Доступные подпапки: {', '.join(subfolders)}" if subfolders else ""
                return f"Error: Folder '{folder_name}' not found in project '{name}'.{hint}"
        else:
            if subfolders:
                return (
                    f"Error: 'folder_name' is required — project '{name}' has subfolders. "
                    f"Укажи подпапку: {', '.join(subfolders)}. "
                    f"Пример: project(action='move_file', name='{name}', folder_name='{subfolders[0]}', file_path='...')"
                )
            target_dir = project_dir

        clean_name = self._clean_telegram_filename(src.name)
        dst = target_dir / clean_name
        shutil.move(str(src), str(dst))
        return f"Файл '{clean_name}' перемещён в проект '{name}/{folder_name}'" if folder_name else f"Файл '{clean_name}' перемещён в проект '{name}'"
